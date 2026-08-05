"""交底书 → 标准专利格式 Word 文档导出

按中国专利说明书格式要求生成 .docx：
- A4 纵向，页边距 上25 / 左25 / 右15 / 下15 mm
- 宋体小四（12pt），1.5 倍行距，纯黑色
- 发明名称第一页第一行居中
- 保留 [0001] 段落编号
- LaTeX 公式（$...$ / $$...$$）经 matplotlib mathtext 渲染为图片插入

使用：
    from src.utils.word_exporter import export_disclosure_to_word
    export_disclosure_to_word(disclosure_text, "output/交底书.docx")
"""

import re
import tempfile
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Pt, Mm, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn

# 页面规格（mm）
PAGE_WIDTH_MM = 210
PAGE_HEIGHT_MM = 297
MARGIN_TOP_MM = 25
MARGIN_LEFT_MM = 25
MARGIN_RIGHT_MM = 15
MARGIN_BOTTOM_MM = 15

# LaTeX 公式：$$...$$ / $...$ / \(...\)，或裸 LaTeX 行
FORMULA_RE = re.compile(
    r"\$\$(.+?)\$\$|\$(.+?)\$|\\\((.+?)\\\)", re.DOTALL)
# 裸 LaTeX 判断：一行含 LaTeX 命令且主要是数学符号
LATEX_CMD_RE = re.compile(r"\\[a-zA-Z]{2,}")

# mathtext 不支持的命令清洗（\left...\right 成对定界 → 普通括号）
_LATEX_CLEAN = [
    (r"\left|", "|"), (r"\right|", "|"),
    (r"\left(", "("), (r"\right)", ")"),
    (r"\left[", "["), (r"\right]", "]"),
    (r"\left\{", "{"), (r"\right\}", "}"),
    (r"\left.", ""), (r"\right.", ""),
    (r"\left", ""), (r"\right", ""),
    (r"\begin{matrix}", ""), (r"\end{matrix}", ""),
    (r"\begin{cases}", ""), (r"\end{cases}", ""),
    (r"\text{", ""),
]


def _looks_like_latex_line(text: str) -> bool:
    """判断一行是否为裸 LaTeX 公式（无 $ 包裹）

    需含 LaTeX 命令，且剔除命令后剩余字符中中文占比低（避免把
    "其中，\\(\\dot\\varepsilon\\)为..." 这种含公式的中文句误判为整行公式）。
    """
    t = text.strip()
    if not t:
        return False
    if not LATEX_CMD_RE.search(t):
        return False
    # 剔除 LaTeX 命令后看剩余字符
    no_cmd = LATEX_CMD_RE.sub("", t)
    no_cmd = re.sub(r"[^一-鿿A-Za-z0-9()=+\-*/.,]", "", no_cmd)
    if not no_cmd:
        return True  # 纯公式
    chinese = sum(1 for c in no_cmd if "一" <= c <= "鿿")
    return chinese / len(no_cmd) < 0.3


def _clean_latex(latex: str) -> str:
    """清洗 mathtext 不支持的 LaTeX 命令（\left/\right、环境、\text）"""
    for old, new in _LATEX_CLEAN:
        latex = latex.replace(old, new)
    # \text{...} 移除后清理残留的孤立花括号
    latex = latex.replace("}", "") if latex.count("{") > latex.count("}") else latex
    return latex.strip()


def _set_run_font(run, name_cn: str = "宋体", size_pt: int = 12,
                  bold: bool = False):
    """设置 run 字体（含中文字体）"""
    run.font.name = "Times New Roman"
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)
    # 设置中文字体
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), name_cn)


def _render_formula_png(latex: str, fontsize: int = 14, dpi: int = 200) -> Optional[Path]:
    """将 LaTeX 公式渲染为 PNG（matplotlib mathtext，无需完整 LaTeX）

    Args:
        latex: LaTeX 公式（不含 $ 定界符）
        fontsize: 字号
        dpi: 分辨率

    Returns:
        PNG 临时文件路径，失败返回 None
    """
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    # 清洗：转义反斜杠 + mathtext 不支持的命令（\left\right等）
    latex = _clean_latex(latex.replace("\\\\", "\\").strip())

    try:
        fig = plt.figure(figsize=(0.1, 0.1))
        t = fig.text(0.5, 0.5, f"${latex}$",
                     fontsize=fontsize, ha="center", va="center")
        fig.canvas.draw()  # 触发渲染，提前暴露 LaTeX 错误

        tmp = Path(tempfile.gettempdir()) / "patent_formula.png"
        fig.savefig(tmp, dpi=dpi, bbox_inches="tight",
                    pad_inches=0.03, transparent=True)
        plt.close(fig)
        return tmp if tmp.exists() else None
    except Exception as e:
        plt.close("all")
        print(f"  公式渲染失败（回退为文本）: {e}")
        return None


def export_disclosure_to_word(disclosure: str, output_path: str,
                              title: Optional[str] = None) -> str:
    """导出交底书为标准专利格式 Word

    Args:
        disclosure: 交底书全文（支持 Markdown 标题和 $...$ LaTeX 公式）
        output_path: 输出 .docx 路径
        title: 发明名称（缺省从全文提取第一个 # 标题）

    Returns:
        保存的路径
    """
    doc = Document()

    # ── 页面设置 ──
    sec = doc.sections[0]
    sec.page_width = Mm(PAGE_WIDTH_MM)
    sec.page_height = Mm(PAGE_HEIGHT_MM)
    sec.top_margin = Mm(MARGIN_TOP_MM)
    sec.left_margin = Mm(MARGIN_LEFT_MM)
    sec.right_margin = Mm(MARGIN_RIGHT_MM)
    sec.bottom_margin = Mm(MARGIN_BOTTOM_MM)

    # Normal 样式：宋体 12pt
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5

    # ── 提取发明名称：优先标准格式 "## 发明名称" 章节的下一行 ──
    m = re.search(r"##\s*发明名称\s*\n\s*([^\n]+)", disclosure)
    if m:
        title = m.group(1).strip()
    elif not title or title in ("专利交底书", "技术交底书"):
        m2 = re.search(r"^#\s+(.+)$", disclosure, re.M)
        title = m2.group(1).strip() if m2 else "专利交底书"

    # 名称第一页第一行，居中
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_title.add_run(title)
    _set_run_font(run, size_pt=14, bold=True)
    doc.add_paragraph()  # 名称与正文之间空一行

    # ── 按行处理正文 ──
    # 移除"发明名称"章节（顶部已作为居中标题输出，避免重复）
    disclosure_clean = re.sub(
        r"##\s*发明名称\s*\n\s*[^\n]+\n?", "", disclosure, count=1)
    lines = disclosure_clean.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # 跳过文档一级标题行（已作为名称输出）
        if stripped.startswith("# ") and i == 0:
            i += 1
            continue

        # Markdown 标题 → 加粗段落
        if stripped.startswith("## "):
            heading = stripped[3:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(heading)
            _set_run_font(run, size_pt=13, bold=True)
            i += 1
            continue

        # 跳过裸 LaTeX 定界符行（\[ \] 成对包裹的显示公式）
        if stripped in (r"\[", r"\]"):
            i += 1
            continue

        # 公式段：$$...$$ 或 $...$ 或裸 LaTeX 行 → 渲染为图片
        is_formula = (stripped.startswith("$$")
                      or (stripped.startswith("$") and stripped.endswith("$")
                          and len(stripped) > 6)
                      or _looks_like_latex_line(stripped))
        if is_formula:
            latex = stripped.strip("$").strip()
            img = _render_formula_png(latex)
            if img:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                from docx.shared import Inches
                p.add_run().add_picture(str(img), width=Inches(3.5))
            else:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(stripped)
                _set_run_font(run, bold=True)
            i += 1
            continue

        # 空行 → 跳过（用段落间距控制）
        if not stripped:
            i += 1
            continue

        # 普通文本段：行内 $...$ 公式拆出渲染
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0.74)  # 两字符缩进
        rest = line
        while rest:
            fm = FORMULA_RE.search(rest)
            if fm:
                # 公式前的文本
                before = rest[:fm.start()]
                if before.strip():
                    run = p.add_run(before)
                    _set_run_font(run)
                # 渲染公式
                latex = (fm.group(1) or fm.group(2) or fm.group(3))
                img = _render_formula_png(latex)
                if img:
                    from docx.shared import Inches
                    p.add_run().add_picture(str(img), width=Inches(2.8))
                else:
                    run = p.add_run(f"${latex}$")
                    _set_run_font(run)
                rest = rest[fm.end():]
            else:
                if rest.strip():
                    run = p.add_run(rest)
                    _set_run_font(run)
                rest = ""
        i += 1

    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return str(out)


if __name__ == "__main__":
    import sys
    if len(sys.argv) < 2:
        print("用法: python src/utils/word_exporter.py <交底书.md> [输出.docx]")
        sys.exit(1)
    src = Path(sys.argv[1])
    text = src.read_text(encoding="utf-8")
    out = sys.argv[2] if len(sys.argv) > 2 else str(src.with_suffix(".docx"))
    path = export_disclosure_to_word(text, out)
    print(f"已导出: {path}")
